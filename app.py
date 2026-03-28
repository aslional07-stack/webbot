from flask import Flask, request, jsonify, send_file, render_template
import requests
from bs4 import BeautifulSoup
from groq import Groq
from docx import Document
from urllib.parse import unquote, urlparse, parse_qs

app = Flask(__name__)

def gercek_url_al(duckduckgo_url):
    try:
        parsed = urlparse(duckduckgo_url)
        params = parse_qs(parsed.query)
        if "uddg" in params:
            return unquote(params["uddg"][0])
    except:
        pass
    return duckduckgo_url

def web_tara(sorgu):
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36"
    }

    url = f"https://html.duckduckgo.com/html/?q={sorgu}"
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, "html.parser")

    urls = []
    for a in soup.find_all("a", class_="result__a")[:3]:
        gercek = gercek_url_al(a["href"])
        urls.append(gercek)

    tum_icerik = ""
    for url in urls:
        try:
            r = requests.get(url, headers=headers, timeout=8)
            s = BeautifulSoup(r.text, "html.parser")
            for p in s.find_all("p")[:15]:
                if p.text and len(p.text) > 30:
                    tum_icerik += p.text + "\n"
        except:
            pass

    return urls, tum_icerik

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/tara", methods=["POST"])
def tara():
    data = request.json
    api_key = data["api_key"]
    sorgu = data["sorgu"]

    urls, tum_icerik = web_tara(sorgu)

    if not tum_icerik:
        tum_icerik = "Sitelerden icerik alinamadi."

    groq_client = Groq(api_key=api_key)
    cevap = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "user",
                "content": f"Asagidaki metni Turkce olarak ozetle ve ana noktalari listele:\n\n{tum_icerik[:3000]}"
            }
        ]
    )
    ozet = cevap.choices[0].message.content

    doc = Document()
    doc.add_heading(f"Rapor: {sorgu}", 0)
    doc.add_heading("Ozet", 1)
    doc.add_paragraph(ozet)
    doc.add_heading("Kaynaklar", 1)
    for url in urls:
        doc.add_paragraph(url)

    doc.save("rapor.docx")
    return jsonify({"durum": "tamam"})

@app.route("/indir")
def indir():
    return send_file("rapor.docx", as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
